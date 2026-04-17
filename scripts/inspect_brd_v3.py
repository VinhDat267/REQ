"""Inspect Group1_Tutorial01_BRD_Ver3.docx to map its current structure."""
import sys
import io
from pathlib import Path
from docx import Document

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

PATH = Path(r"c:/Users/VinhDat/Desktop/REQ/Group1_Tutorial01_BRD_Ver3.docx")

doc = Document(str(PATH))

print(f"=== DOC: {PATH.name} ===")
print(f"Paragraphs: {len(doc.paragraphs)}  Tables: {len(doc.tables)}  Sections: {len(doc.sections)}")
print()

# Dump all paragraphs with style name and first 120 chars
print("=== PARAGRAPH DUMP ===")
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else "None"
    txt = (p.text or "").strip().replace("\n", " / ")
    if not txt and "Heading" not in style:
        continue
    if len(txt) > 120:
        txt = txt[:117] + "..."
    print(f"[{i:3d}] <{style}> {txt}")

print()
print("=== TABLE SUMMARY ===")
for ti, t in enumerate(doc.tables):
    print(f"--- TABLE {ti}: rows={len(t.rows)} cols={len(t.columns)}")
    # print up to first 2 rows preview
    for ri, row in enumerate(t.rows[:3]):
        cells_preview = [c.text.strip().replace("\n", " / ")[:60] for c in row.cells]
        print(f"   row{ri}: {cells_preview}")
    if len(t.rows) > 3:
        print(f"   ... +{len(t.rows)-3} more rows")
