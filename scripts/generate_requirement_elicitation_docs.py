from copy import deepcopy
from pathlib import Path
from shutil import copyfile

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_DIR = ROOT / "RequirementElicitationTemplates"
OUTPUT_DIR = ROOT / "docs" / "requirement-elicitation"


TEAM_MEMBERS = [
    "Nguyen Thi Hai My (Business Analyst)",
    "Nguyen Dat Vinh (Project Lead / System Analyst)",
    "Vuong Gia Ly (Business Analyst)",
    "Nguyen Minh Duc (System Analyst)",
    "Dang Khanh Huyen (Business Analyst)",
    "Le Thanh Dat (System Analyst)",
    "Nguyen Dinh Chien (QA / Documentation)",
    "Nguyen Thien Hieu (System Analyst)",
]

SPONSOR = "Nguyen Quang Anh"
SPONSOR_TITLE = "CEO / Project Sponsor"
SPONSOR_EMAIL = "N/A - contact detail not provided in current project documents"
PROJECT_NAME = "Wonton POS"
RESTAURANT_NAME = "HongKong A Hai Crispy Roast Duck Restaurant"


def fill_cell(table, row_idx, cell_idx, text):
    row = table.rows[row_idx]
    if cell_idx < len(row.cells):
        row.cells[cell_idx].text = text


def fill_paragraph(doc, idx, text):
    if idx < len(doc.paragraphs):
        doc.paragraphs[idx].text = text


def clear_paragraph(doc, idx):
    if idx >= len(doc.paragraphs):
        return
    paragraph = doc.paragraphs[idx]
    element = paragraph._element
    for child in list(element):
        if child.tag.endswith("}pPr"):
            continue
        element.remove(child)


def clone_table_before_paragraph(doc, source_table_idx, paragraph_idx):
    new_table_element = deepcopy(doc.tables[source_table_idx]._element)
    doc.paragraphs[paragraph_idx]._element.addprevious(new_table_element)
    return doc.tables[-1]


def ensure_question_tables(doc, first_question_table_idx, paragraph_idx, total_questions):
    current_question_count = len(doc.tables) - first_question_table_idx
    while current_question_count < total_questions:
        clone_table_before_paragraph(doc, first_question_table_idx, paragraph_idx)
        current_question_count += 1


def make_output(template_name, output_name, editor):
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / output_name
    copyfile(TEMPLATE_DIR / template_name, output_path)
    doc = Document(output_path)
    editor(doc)
    doc.save(output_path)
    return output_path


def edit_prepare(doc):
    fill_paragraph(
        doc,
        2,
        "Elicit business and user requirements for Wonton POS, focusing on online ordering, "
        "counter operations, payment handling, pickup scheduling, order tracking, table "
        "management, kitchen coordination, staff administration, and revenue reporting within "
        "the current project scope.",
    )
    fill_paragraph(
        doc,
        6,
        "Primary participants: Nguyen Quang Anh (project sponsor), operational representatives "
        "from cashier/service/kitchen workflows, and the project team: "
        + ", ".join(TEAM_MEMBERS)
        + ".",
    )
    fill_paragraph(
        doc,
        8,
        "Language: Vietnamese for working sessions and English for finalized project artifacts. "
        "Expertise: the sponsor provides business goals and policy decisions; operational "
        "representatives validate the as-is workflow and pain points; the BA/SA team converts "
        "findings into BRD content, business rules, and use cases. Location: onsite sessions at "
        f"{RESTAURANT_NAME} plus follow-up confirmation through direct communication channels. "
        "Influence: the sponsor validates scope and priorities, while operational representatives "
        "validate feasibility and usability.",
    )
    fill_paragraph(
        doc,
        10,
        "Determine logistics - Session set: stakeholder interview (60 minutes), workflow "
        "observation (75 minutes during operating hours), questionnaire collection window "
        "(3 days), and findings-validation meeting (75 minutes). Participants: sponsor, "
        "operational representatives, and BA/SA team members. Resources: elicitation templates, "
        "BRD draft, business rules draft, menu references, laptop, mobile devices, and "
        "note-taking materials.",
    )
    clear_paragraph(doc, 11)
    fill_paragraph(
        doc,
        15,
        "Produce completed elicitation records using the provided templates, a validated scope "
        "statement, clarified business rules, prioritized use cases, and documented assumptions, "
        "constraints, issues, and risks for the Wonton POS project.",
    )
    fill_paragraph(
        doc,
        19,
        "Supporting materials prepared for elicitation: BRD draft, Business Rules draft, current "
        "menu references, use case list, glossary, order and payment rule summary, and process "
        "notes for dine-in, takeaway, and pickup workflows.",
    )
    fill_paragraph(
        doc,
        23,
        "Use the provided DOCX templates as the primary elicitation record set. Supporting raw "
        "notes, form links, and working summaries may be captured separately, but the finalized "
        "project record is maintained in these completed templates.",
    )
    fill_paragraph(
        doc,
        28,
        "Stakeholders receive an invitation or direct briefing that covers the session topic, "
        "purpose, expected duration, location or channel, required preparation, and the expected "
        "output of each elicitation activity.",
    )
    fill_paragraph(
        doc,
        31,
        "Key risks: limited stakeholder availability, incomplete operational detail, observation "
        "noise during busy hours, and inconsistent interpretation across roles. Mitigation: "
        "triangulate findings across interview, questionnaire, observation, and meeting notes; "
        "recap decisions at the end of each session; and keep assumptions and open issues visible "
        "for follow-up confirmation.",
    )


def edit_interview(doc):
    fill_cell(
        doc.tables[0],
        0,
        1,
        "Capture business goals, operational pain points, scope boundaries, business rules, and "
        "success expectations for Wonton POS from the primary business stakeholder.",
    )
    fill_cell(doc.tables[0], 1, 1, SPONSOR)
    fill_cell(doc.tables[1], 0, 1, SPONSOR_TITLE)
    fill_cell(doc.tables[1], 0, 3, SPONSOR_EMAIL)
    fill_cell(doc.tables[1], 1, 1, "06/03/2026")
    fill_cell(doc.tables[1], 1, 3, "09:00 AM - 10:15 AM")
    fill_cell(doc.tables[1], 2, 1, f"Onsite discussion area at {RESTAURANT_NAME}")
    fill_cell(doc.tables[1], 2, 3, "Written notes with post-session recap confirmation")

    questions = [
        (
            "What are the biggest operational problems in the current ordering process?",
            "Manual order capture causes missed items, repeated clarification between counter staff "
            "and kitchen staff, and slow handoff during busy periods.",
        ),
        (
            "Which service models must the system support?",
            "Dine-in, Takeaway, and Pickup must be supported in the current scope. Delivery "
            "integration remains outside the current scope.",
        ),
        (
            "What payment rule should apply to each service model?",
            "Takeaway and Pickup must be prepaid before kitchen processing. Dine-in orders may "
            "remain unpaid until counter settlement after service.",
        ),
        (
            "What customer information is required when someone orders without creating an account?",
            "Guest checkout should work with at least customer name and phone number so the order "
            "can be created quickly without full account registration.",
        ),
        (
            "How should guest customers track their orders after checkout?",
            "Guest customers should track orders using Order Code and Phone Number instead of "
            "needing an account.",
        ),
        (
            "What menu details and customization options must be available to customers?",
            "Customers need to view menu items, prices, availability, toppings, quantities, and "
            "special notes before placing an order.",
        ),
        (
            "How should sold-out or temporarily unavailable items be handled?",
            "Managers should be able to mark items unavailable in the system immediately so both "
            "staff and customers stop selecting them.",
        ),
        (
            "What information must the kitchen receive for every order?",
            "The kitchen needs item name, quantity, toppings, notes, service model, pickup time if "
            "applicable, and the current operational status of the order.",
        ),
        (
            "How should dine-in table management work in practice?",
            "Managers should maintain table data, while FOH staff need a fast view of table status "
            "and the ability to assign or transfer dine-in orders to the correct table.",
        ),
        (
            "What staff roles or permissions are essential in the system?",
            "The system needs manager, cashier, server, and kitchen-oriented responsibilities, and "
            "staff accounts should support multi-role assignment where needed.",
        ),
        (
            "How should the counter order flow work for walk-in customers?",
            "FOH staff should be able to create dine-in and takeaway orders quickly at the counter, "
            "apply the correct payment rule, and push valid orders into the next workflow step.",
        ),
        (
            "What rules should control pickup scheduling?",
            "Pickup orders should let customers choose a valid time slot, and the system should "
            "prevent overbooking based on available service capacity.",
        ),
        (
            "Which payment methods should be recognized in the system?",
            "Cash, QR payment, and supported online payment flows are needed, with sandbox or test "
            "integration during the current implementation phase.",
        ),
        (
            "What kinds of status notifications are important for customers or staff?",
            "Status changes such as order accepted, cooking, ready, and completed should be visible "
            "quickly so customers and staff can coordinate without repeated verbal checks.",
        ),
        (
            "What reports or analytics are most valuable to the business owner?",
            "Revenue statistics, payment-method breakdown, top-selling items, peak-hour patterns, "
            "and order volume by service model are the most valuable management views.",
        ),
        (
            "Are there compliance, technology, or deployment constraints we must respect?",
            "The project should stay web-based, remain simple to operate, use test environments for "
            "payment integration, and avoid unnecessary hardware or deployment complexity.",
        ),
        (
            "Which features should remain out of scope for now?",
            "Delivery platform integration, promotions, loyalty programs, payroll, and native "
            "mobile apps should remain outside the current project scope.",
        ),
        (
            "What would make Wonton POS a successful project from your perspective?",
            "Success means fewer order mistakes, faster coordination across roles, clear order and "
            "payment visibility, easier menu and staff maintenance, and trustworthy revenue data "
            "for decision-making.",
        ),
    ]

    ensure_question_tables(doc, first_question_table_idx=2, paragraph_idx=19, total_questions=len(questions))
    fill_paragraph(doc, 5, f"Interview Questions ({len(questions)} questions)")

    for idx, (question, note) in enumerate(questions, start=2):
        fill_cell(doc.tables[idx], 0, 0, f"Question {idx - 1}:")
        fill_cell(doc.tables[idx], 0, 1, question)
        fill_cell(doc.tables[idx], 1, 0, ":")
        fill_cell(doc.tables[idx], 1, 1, note)

    fill_paragraph(
        doc,
        20,
        "Detailed interview notes:\n"
        "1. Business priority: the sponsor emphasized that the first business value of the system "
        "is not visual complexity but operational reliability. The restaurant needs fewer order "
        "mistakes, faster order handoff, and clearer visibility across cashier, service, and "
        "kitchen roles.\n"
        "2. Service-model clarification: dine-in, takeaway, and pickup are all required because "
        "they reflect real current operations. Delivery-related capability was explicitly treated "
        "as outside the present requirement baseline to avoid scope expansion.\n"
        "3. Payment-policy clarification: takeaway and pickup must not proceed to kitchen handling "
        "before payment is confirmed, while dine-in may remain unpaid until later settlement. This "
        "means the system must preserve independent order and payment states instead of forcing a "
        "single lifecycle.\n"
        "4. Customer experience expectation: the sponsor prefers a simple Vietnamese-first "
        "interface with minimal training effort for staff and low friction for customers. Guest "
        "checkout remains important because many customers may place one-off orders without wanting "
        "to register an account.\n"
        "5. Tracking and transparency: the sponsor confirmed that guest order tracking must work "
        "through Order Code + Phone Number. This was described as essential for reducing repeated "
        "status inquiries and helping customers trust the ordering process.\n"
        "6. Reporting expectation: management views should support revenue tracking, top-selling "
        "items, and peak-hour insight because the restaurant currently lacks dependable structured "
        "data for these decisions.\n"
        "7. Technical and rollout constraint: online payment can use sandbox or test integration "
        "during implementation, but the business rule around prepayment must still behave as if it "
        "were a real operational flow.\n"
        "8. Follow-up implication: requirement analysis must pay special attention to order-status "
        "visibility, table assignment, item availability handling, and fast staff workflows at the "
        "counter.",
    )


def edit_observation(doc):
    fill_cell(
        doc.tables[0],
        0,
        1,
        "Observe the as-is dine-in and takeaway workflow to identify manual handoff issues, order "
        "status gaps, payment checks, table usage pain points, and reporting inefficiencies.",
    )
    fill_cell(
        doc.tables[0],
        1,
        1,
        "Operational staff representatives (cashier, service staff, kitchen staff)",
    )
    fill_cell(doc.tables[1], 0, 1, "Observed operational roles (anonymized in this summary)")
    fill_cell(doc.tables[1], 0, 3, "N/A - observation participants summarized by role")
    fill_cell(doc.tables[1], 1, 1, "06/03/2026")
    fill_cell(doc.tables[1], 1, 3, "11:00 AM - 12:15 PM")
    fill_cell(doc.tables[1], 2, 1, "Passive / non-participant observation")
    fill_cell(doc.tables[1], 2, 3, f"Cashier counter and kitchen pass area at {RESTAURANT_NAME}")

    fill_paragraph(
        doc,
        5,
        "Detailed observation notes:\n"
        "1. Observation context: the workflow was reviewed from the cashier counter toward the "
        "kitchen handoff area during a real operating period so the team could see where manual "
        "coordination created delay or confusion.\n"
        "2. Order capture behavior: counter staff often receive orders verbally or on paper, then "
        "restate them to the kitchen. This creates repeated handoff, duplicated effort, and more "
        "opportunity for item, quantity, or note mismatch.\n"
        "3. Availability handling: sold-out cases appear to rely on informal verbal confirmation "
        "with the kitchen rather than immediate shared system visibility. This slows order "
        "finalization and can produce inconsistent information for customers.\n"
        "4. Table tracking behavior: dine-in table assignment is handled manually, which makes it "
        "harder to know which order is linked to which table during busy periods or when staff "
        "switch tasks.\n"
        "5. Status-visibility gap: takeaway and pickup customers frequently need manual status "
        "updates because there is no obvious shared real-time order state visible across staff and "
        "customers.\n"
        "6. Payment-control gap: prepaid order confirmation depends on manual checking instead of "
        "system-enforced validation before the order continues to kitchen processing.\n"
        "7. Reporting gap: end-of-day revenue recap still depends heavily on manual counting, "
        "which increases the risk of missing transactions, inconsistent totals, or delayed review.\n"
        "8. Requirement implication: the project needs clearer digital order capture, fast item "
        "availability control, explicit order and payment status visibility, better table mapping, "
        "workflow support for prepaid enforcement, and more reliable automated reporting.\n"
        "9. Validation note: the observation findings are strongest where they match sponsor "
        "interview statements and questionnaire feedback, especially around handoff problems, "
        "prepaid enforcement, and lack of reporting confidence.",
    )


def edit_questionnaire(doc):
    fill_paragraph(doc, 0, "Questionnaire Notes:")
    fill_paragraph(doc, 1, "\tQuestionnaire Details\t")
    fill_paragraph(doc, 5, "Questionnaire Questions")
    fill_paragraph(doc, 36, "Questionnaire Notes")
    fill_paragraph(doc, 49, "Questionnaire Outcomes/Learnings")

    fill_cell(doc.tables[0], 0, 0, "Purpose of Questionnaire:")
    fill_cell(
        doc.tables[0],
        0,
        1,
        "Collect structured feedback about operational pain points, priority features, workflow "
        "rules, usability needs, and rollout constraints for Wonton POS.",
    )
    fill_cell(
        doc.tables[0],
        1,
        1,
        "Project sponsor, manager or owner representative, cashier representative, service staff "
        "representative, kitchen representative, and BA or SA team members consolidating feedback",
    )
    fill_cell(doc.tables[1], 0, 1, "Google Forms link shared directly with participants")
    fill_cell(doc.tables[1], 1, 1, "07/03/2026")
    fill_cell(doc.tables[1], 1, 3, "09/03/2026")
    fill_cell(
        doc.tables[2],
        0,
        1,
        "No monetary incentive. Participants were informed that the questionnaire would be used "
        "to clarify workflow needs, reduce ambiguity, and strengthen the project requirement set.",
    )

    questions = [
        ("Which service model currently creates the most operational pressure?", "Multiple choice", "Dine-in; Takeaway; Pickup; All are equally difficult", "Identify which service flow needs the most design attention."),
        ("How often do order mistakes happen during busy periods?", "Likert scale (Never to Very Often)", "Never; Rarely; Sometimes; Often; Very Often", "Measure the perceived severity of order accuracy issues."),
        ("Which part of the current process wastes the most time?", "Multiple choice", "Taking orders; Relaying orders to kitchen; Tracking order status; Payment settlement; End-of-day reporting", "Locate the highest-friction operational step."),
        ("What should trigger kitchen preparation for takeaway and pickup orders?", "Single choice", "Immediately after order creation; Only after payment confirmation; Manual staff confirmation; Other", "Validate the prepaid processing rule for applicable service models."),
        ("How important is guest checkout for online customers?", "Likert scale", "Not important; Slightly important; Moderately important; Important; Critical", "Assess whether account-less ordering is necessary for adoption."),
        ("What information should guest checkout require at minimum?", "Checkbox", "Full name; Phone number; Email; Pickup note; No extra information beyond item selection", "Confirm the minimum data set for fast guest ordering."),
        ("How should customers track their orders if they do not create an account?", "Single choice", "Order code only; Phone number only; Order code + phone number; Contact staff directly", "Confirm the preferred guest order tracking mechanism."),
        ("What order information must always be visible to the kitchen?", "Checkbox", "Menu item; Quantity; Toppings; Notes; Service model; Pickup time; Payment status", "Define the minimum kitchen-display information set."),
        ("How often does the menu change in a typical week?", "Multiple choice", "Never; 1-2 times; 3-5 times; Daily", "Understand how frequently menu maintenance must be performed."),
        ("How should sold-out items be handled during live operations?", "Single choice", "Manual verbal notice only; Mark unavailable in system; Hide from customer view only; Kitchen decides case by case", "Define the system response to menu availability changes."),
        ("Which staff management capabilities are essential?", "Checkbox", "Create account; Update account; Deactivate account; Assign multiple roles; View staff list", "Prioritize admin functionality for employee management."),
        ("How important is multi-role support for employee accounts?", "Likert scale", "Not important; Slightly important; Moderately important; Important; Critical", "Validate the need for staff accounts that hold more than one role."),
        ("How should tables be represented for dine-in operations?", "Single choice", "Only table number; Table number + status; Layout view + status; No table tracking needed", "Confirm the required depth of table management."),
        ("When should a table be marked as occupied or available again?", "Multiple choice", "Occupied at order creation; Occupied at food service; Available after payment; Available after staff manual reset", "Clarify operational rules for table-state updates."),
        ("Which payment methods should the system recognize?", "Checkbox", "Cash; QR payment; Online sandbox gateway; Card payment; Split payment not needed", "Clarify supported payment scenarios for the current scope."),
        ("How important is separating order status from payment status?", "Likert scale", "Not important; Slightly important; Moderately important; Important; Critical", "Validate the need for distinct operational and financial lifecycles."),
        ("Which step causes the most delay in counter payment processing?", "Multiple choice", "Confirming amount; Selecting payment method; Waiting for payment confirmation; Printing or sending receipts; Marking order complete", "Pinpoint where counter payment flow should be simplified."),
        ("Which dashboard metrics are most useful to the business owner?", "Checkbox", "Daily revenue; Weekly revenue; Top-selling items; Peak hours; Payment method split; Order volume by service model", "Prioritize analytics outputs for the reporting dashboard."),
        ("Which report filters are important when reviewing performance?", "Checkbox", "Date range; Service model; Payment method; Menu category; Staff shift", "Clarify reporting flexibility required by management."),
        ("What response speed feels acceptable for common system actions?", "Single choice", "Under 1 second; Under 2 seconds; Under 3 seconds; More than 3 seconds is acceptable", "Capture usability expectations for system performance."),
        ("Which devices are most likely to be used by staff during daily operations?", "Checkbox", "Desktop PC; Laptop; Tablet; Mobile phone", "Guide interface optimization and device support."),
        ("What is the most realistic deployment environment for the system?", "Single choice", "Desktop browser only; Mobile browser only; Responsive web on both desktop and mobile; Native mobile app", "Confirm the platform strategy for the project."),
        ("How much training should staff need before using the system confidently?", "Single choice", "No training; Under 30 minutes; 30-60 minutes; More than 1 hour", "Set expectations for onboarding and usability."),
        ("Which status notifications are most useful to customers?", "Checkbox", "Order accepted; Cooking started; Ready for pickup or serving; Completed; Cancelled", "Clarify the notification events that provide the most value."),
        ("How should pickup capacity be controlled?", "Single choice", "No capacity limit; Fixed time slots; Staff approval only; Daily manual list", "Define the pickup scheduling mechanism."),
        ("How important is automatic cancellation for unpaid prepaid orders after a timeout?", "Likert scale", "Not important; Slightly important; Moderately important; Important; Critical", "Validate timeout handling for prepayment-required flows."),
        ("How valuable is the order rating feature for completed orders?", "Single choice", "Not needed; Nice to have; Important for service improvement; Critical for customer feedback tracking", "Assess business value of post-order rating capability."),
        ("Which features should definitely stay out of the current scope?", "Checkbox", "Delivery integration; Promotions; Loyalty program; Payroll; Native mobile app", "Reconfirm boundaries and prevent scope creep."),
        ("What is the biggest risk if the project requirements are misunderstood?", "Short answer", "Open text", "Capture hidden concerns that may not appear in structured questions."),
        ("Which outcome would make the system clearly valuable to the restaurant?", "Short answer", "Open text", "Identify the clearest business-value criteria for solution prioritization."),
    ]

    ensure_question_tables(doc, first_question_table_idx=3, paragraph_idx=36, total_questions=len(questions))
    fill_paragraph(doc, 5, f"Questionnaire Questions ({len(questions)} questions)")

    for idx, (question, answer_type, answers, purpose) in enumerate(questions, start=3):
        question_number = idx - 2
        fill_cell(doc.tables[idx], 0, 0, f"Question {question_number}:")
        fill_cell(doc.tables[idx], 0, 1, question)
        fill_cell(doc.tables[idx], 1, 0, "Answer Type:")
        fill_cell(doc.tables[idx], 1, 1, answer_type)
        fill_cell(doc.tables[idx], 2, 0, "Answers:")
        fill_cell(doc.tables[idx], 2, 1, answers)
        fill_cell(doc.tables[idx], 3, 0, "Question's Purpose:")
        fill_cell(doc.tables[idx], 3, 1, purpose)

    fill_paragraph(
        doc,
        37,
        "Detailed questionnaire notes:\n"
        "1. Questionnaire structure: the question set was intentionally grouped into customer "
        "ordering, counter workflow, kitchen workflow, payment handling, reporting, usability, and "
        "scope-control themes so that responses could be compared across operational viewpoints.\n"
        "2. Response interpretation approach: questions combined multiple-choice, checkbox, Likert, "
        "and short-answer formats in order to capture both measurable preference patterns and "
        "qualitative concerns that might not surface in a simple interview summary.\n"
        "3. Validation use: the questionnaire was not treated as an isolated source. Instead, its "
        "purpose was to confirm, challenge, or refine findings already emerging from interview and "
        "observation activities.\n"
        "4. Conflict handling: if any answer suggested a broader feature direction, that response "
        "was checked against current scope, operational reality, and sponsor intent before being "
        "treated as a requirement candidate.",
    )
    fill_paragraph(
        doc,
        38,
        "Detailed response-consolidation notes:\n"
        "1. Role-based feedback was consolidated at summary level to keep the elicitation record "
        "usable while avoiding overexposure of individual respondents in the formal project "
        "artifact.\n"
        "2. Repeated response themes were treated as stronger evidence than isolated suggestions, "
        "especially where they matched observable workflow pain points such as delayed kitchen "
        "handoff, unclear status visibility, or manual reporting effort.\n"
        "3. Responses that pointed to device support, training needs, and interface simplicity were "
        "used to inform non-functional expectations and rollout assumptions, not only functional "
        "requirements.",
    )
    fill_paragraph(
        doc,
        39,
        "Detailed scope-filtering notes:\n"
        "1. Suggestions outside the current scope, such as delivery platform integration, loyalty, "
        "payroll, and native mobile apps, were logged as future considerations rather than current "
        "requirements.\n"
        "2. This filtering step was necessary to keep the requirements baseline aligned with sponsor "
        "intent and to avoid mixing operational priorities with later-stage expansion ideas.\n"
        "3. Where a response touched both current and future scope, only the immediately actionable "
        "part was retained in the active requirement interpretation.",
    )
    fill_paragraph(
        doc,
        50,
        "Detailed outcomes and learnings:\n"
        "1. The questionnaire reinforced the need for accurate order capture, clear prepaid rules "
        "for takeaway and pickup, simple menu and staff administration, clear table visibility, "
        "fast status updates, and reliable revenue reporting.\n"
        "2. Answers related to guest checkout and order tracking supported the requirement that "
        "non-registered customers must still be able to complete orders and later retrieve status "
        "through Order Code plus Phone Number.\n"
        "3. Multiple questions around payment and workflow sequencing supported the decision to keep "
        "order status and payment status separate, especially for dine-in versus prepaid flows.\n"
        "4. Responses about devices, training time, and interface clarity suggested that the "
        "solution should prioritize responsive web access and low training overhead rather than "
        "complex feature density.\n"
        "5. Reporting-related answers consistently pointed toward management value in revenue, "
        "payment-method split, top-selling items, and peak-hour insight, confirming that analytics "
        "is not optional background functionality but a real stakeholder need.\n"
        "6. Overall, the questionnaire increased confidence that the selected requirement direction "
        "matches both business priorities and operational reality rather than being based on a "
        "single viewpoint alone.",
    )


def edit_meeting(doc):
    fill_paragraph(doc, 0, "Meeting Notes:")
    fill_paragraph(doc, 1, "\tMeeting Details\t")
    fill_paragraph(doc, 4, "Meeting Agenda")
    fill_paragraph(doc, 9, "Meeting Notes")
    fill_paragraph(doc, 19, "Meeting Action Items")

    fill_cell(doc.tables[0], 0, 0, "Purpose of Meeting:")
    fill_cell(
        doc.tables[0],
        0,
        1,
        "Validate elicitation findings, resolve open questions, and agree on the current scope, "
        "priority requirements, and next actions for Wonton POS.",
    )
    fill_cell(doc.tables[1], 0, 1, "10/03/2026")
    fill_cell(doc.tables[1], 0, 3, f"Onsite meeting area at {RESTAURANT_NAME} + online follow-up confirmation")
    fill_cell(doc.tables[1], 1, 1, "Nguyen Dat Vinh")
    fill_cell(doc.tables[1], 1, 3, "Nguyen Dinh Chien")
    fill_cell(doc.tables[1], 2, 1, f"{SPONSOR} (Yes)")
    fill_cell(
        doc.tables[2],
        0,
        1,
        "Operational representatives from cashier/service/kitchen workflows, plus project team: "
        + ", ".join(TEAM_MEMBERS),
    )

    agenda = [
        "Review project goals and the main pain points identified during elicitation.",
        "Confirm business expectations and decision criteria with the sponsor.",
        "Validate in-scope service models: Dine-in, Takeaway, and Pickup.",
        "Confirm actor list and role boundaries for Admin WebApp and Client WebApp.",
        "Review findings from interview and observation activities.",
        "Review structured themes from the questionnaire.",
        "Validate payment rules and the separation of order status versus payment status.",
        "Confirm the priority workflow set across customer, FOH, BOH, and manager activities.",
        "Review assumptions, constraints, and risk mitigation actions.",
        "Confirm next deliverables, owners, and approval checkpoints.",
    ]
    for idx, item in enumerate(agenda):
        fill_cell(doc.tables[3], idx, 0, item)

    fill_paragraph(
        doc,
        10,
        "Detailed meeting notes - Summary:\n"
        "1. Participants aligned that the project must directly address real operational problems "
        "already observed in the restaurant rather than adding features with low immediate value.\n"
        "2. The strongest shared concern remained order accuracy and handoff quality between front-"
        "of-house staff and kitchen staff, especially during busy periods.\n"
        "3. The group also confirmed that reporting is a meaningful management need because manual "
        "daily recap does not provide enough confidence or detail for performance review.",
    )
    fill_paragraph(
        doc,
        11,
        "Detailed meeting notes - Scope and rules:\n"
        "1. Dine-in, Takeaway, and Pickup remain in scope because they represent the actual service "
        "models that the restaurant currently needs to manage.\n"
        "2. Delivery integration, promotions, loyalty, payroll, and native mobile apps remain "
        "outside the current scope to protect implementation focus.\n"
        "3. The group confirmed that prepaid enforcement is mandatory for takeaway and pickup, "
        "while dine-in may continue under delayed payment settlement rules.\n"
        "4. This discussion reinforced the need to model order status and payment status "
        "independently in the requirement set.",
    )
    fill_paragraph(
        doc,
        12,
        "Detailed meeting notes - Functional alignment:\n"
        "1. Customer-side priorities include browsing the menu, placing orders with the correct "
        "service model, completing payment where required, and tracking order progress clearly.\n"
        "2. Staff-side priorities include fast counter order creation, payment processing, table "
        "assignment, availability handling, and smoother kitchen coordination.\n"
        "3. Management-side priorities include menu maintenance, staff administration, table "
        "visibility, and revenue analytics.\n"
        "4. Participants agreed that notifications and kitchen status updates are necessary because "
        "they reduce repeated verbal follow-up and improve shared visibility across roles.",
    )
    fill_paragraph(
        doc,
        13,
        "Detailed meeting notes - Open issues and next-step direction:\n"
        "1. The exact online payment integration path can be finalized during solution design as "
        "long as prepaid flow enforcement and payment-state visibility remain intact.\n"
        "2. The team should continue validating edge cases around sold-out items, pickup capacity, "
        "guest tracking, and timing of table-status changes.\n"
        "3. Follow-up documentation should preserve a clear trace from elicitation findings to "
        "business rules, functional requirements, and use-case behavior so later review remains "
        "defensible.",
    )

    action_items = [
        ("Finalize the scope statement and project overview wording", "Nguyen Dat Vinh", "11/03/2026"),
        ("Refine payment and service-model business rules", "Nguyen Thi Hai My", "11/03/2026"),
        ("Consolidate sponsor interview findings into requirement notes", "Vuong Gia Ly", "11/03/2026"),
        ("Update use-case traceability against elicitation findings", "Nguyen Minh Duc", "12/03/2026"),
        ("Summarize questionnaire themes and unresolved items", "Nguyen Dinh Chien", "12/03/2026"),
        ("Review menu management and availability handling requirements", "Dang Khanh Huyen", "12/03/2026"),
        ("Validate table-management and dine-in assignment rules", "Le Thanh Dat", "12/03/2026"),
        ("Confirm guest checkout and order-tracking fields", "Nguyen Dat Vinh", "13/03/2026"),
        ("Review reporting KPI list with sponsor feedback", "Nguyen Thi Hai My", "13/03/2026"),
        ("Clarify notification events and visibility rules", "Nguyen Minh Duc", "13/03/2026"),
        ("Document remaining assumptions, constraints, and risks", "Nguyen Dinh Chien", "14/03/2026"),
        ("Prepare the next requirements review checkpoint", "Vuong Gia Ly", "14/03/2026"),
    ]
    for row_idx, (item, owner, due_date) in enumerate(action_items, start=1):
        fill_cell(doc.tables[4], row_idx, 0, item)
        fill_cell(doc.tables[4], row_idx, 1, owner)
        fill_cell(doc.tables[4], row_idx, 2, due_date)


def main():
    outputs = []
    outputs.append(
        make_output(
            "0305_RESOURCE_PrepareForElicitationTemplate.docx",
            "01_Prepare_for_Elicitation_Wonton_POS.docx",
            edit_prepare,
        )
    )
    outputs.append(
        make_output(
            "0410_RESOURCE_InterviewElicitationNotes_TheBAGuide.docx",
            "02_Interview_Elicitation_Notes_Wonton_POS.docx",
            edit_interview,
        )
    )
    outputs.append(
        make_output(
            "0410_RESOURCE_SurveyElicitationNotes_TheBAGuide.docx",
            "03_Questionnaire_Elicitation_Notes_Wonton_POS.docx",
            edit_questionnaire,
        )
    )
    outputs.append(
        make_output(
            "0410_RESOURCE_WorkshopNotes_TheBAGuide.docx",
            "04_Meeting_Elicitation_Notes_Wonton_POS.docx",
            edit_meeting,
        )
    )
    outputs.append(
        make_output(
            "0410_RESOURCE_ObservationElicitationNotes_TheBAGuide.docx",
            "05_Observation_Elicitation_Notes_Wonton_POS.docx",
            edit_observation,
        )
    )

    for path in outputs:
        print(path)


if __name__ == "__main__":
    main()
