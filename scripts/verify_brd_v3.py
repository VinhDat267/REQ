"""Verify the BRD v3 uplift. Prints key tables after edit."""
import sys, io
from pathlib import Path
from docx import Document

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

doc = Document(str(Path(r"c:/Users/VinhDat/Desktop/REQ/Group1_Tutorial01_BRD_Ver3.docx")))

print(f"Total paragraphs: {len(doc.paragraphs)}  tables: {len(doc.tables)}")
print()

# Version history
print("=== VERSION HISTORY (Table 0) ===")
for ri, row in enumerate(doc.tables[0].rows):
    vals = [c.text.strip().replace('\n', ' / ')[:80] for c in row.cells]
    print(f"  r{ri}: {vals}")

print()
print("=== FR TABLE HEAD+TAIL (Table 9) ===")
fr = doc.tables[9]
print(f"rows: {len(fr.rows)}")
for ri in [0, 1, 4, 7, 16, 17, 18, 19, 20, 25, 26]:
    if ri >= len(fr.rows):
        continue
    row = fr.rows[ri]
    print(f"  r{ri}: {[c.text.strip()[:60] for c in row.cells]}")

print()
print("=== NFR TABLE (Table 10) ===")
nfr = doc.tables[10]
print(f"rows: {len(nfr.rows)}")
for ri, row in enumerate(nfr.rows):
    vals = [c.text.strip().replace('\n', ' / ')[:70] for c in row.cells]
    print(f"  r{ri}: {vals}")

print()
print("=== UC specs (Tables 27..36) first row ID + rows count ===")
for ti in range(27, min(37, len(doc.tables))):
    t = doc.tables[ti]
    id_name = t.rows[0].cells[1].text.strip() if len(t.rows) > 0 else "?"
    print(f"  T{ti}: rows={len(t.rows)}  id={id_name!r}")

print()
print("=== Glossary tail (Table 4 last 11 rows) ===")
gl = doc.tables[4]
for ri in range(max(0, len(gl.rows) - 11), len(gl.rows)):
    row = gl.rows[ri]
    print(f"  r{ri}: {row.cells[0].text.strip()!r:30s} -> {row.cells[1].text.strip()[:80]}")

print()
print("=== Assumptions table (Table 7) ===")
at = doc.tables[7]
for ri, row in enumerate(at.rows):
    vals = [c.text.strip()[:100] for c in row.cells]
    print(f"  r{ri}: {vals}")
