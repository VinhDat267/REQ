"""
Uplift Group1_Tutorial01_BRD_Ver3.docx from midterm 16-UC orientation
to BRD Final v3 with 26 UC lane, aligned with the locked Business Rules.

Edits:
  A. Table 0  Version History — add v3 Final row
  B. Table 4  Glossary — add Reorder, Write-off, Refund Pending, Refunded,
                         Comp, Forfeited, Shift Close, 86'd Item
  C. Table 7  Assumptions/Constraints — add A8 (scope lock 2026-04-17)
  D. Table 9  FR — drop "midterm" from rationales, add FR-17..FR-26
  E. Table 10 NFR — update scope references; add NFR-08 audit; NFR-09 privacy
  F. Tables 11..26 unchanged (UC-01..UC-16 specs)
  G. Insert 10 new UC specification tables for UC-17..UC-26
     (cloned from UC-16 structure to preserve styling)

Save result in-place.
"""
import sys, io
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

SRC = Path(r"c:/Users/VinhDat/Desktop/REQ/Group1_Tutorial01_BRD_Ver3.docx")

doc = Document(str(SRC))


# ---------- helpers ---------------------------------------------------------

def set_cell_text(cell, text):
    """Replace cell text while preserving paragraph style of the first paragraph."""
    # Keep first paragraph, drop the rest
    p0 = cell.paragraphs[0]
    # Clear runs in first paragraph
    for run in list(p0.runs):
        run._element.getparent().remove(run._element)
    # Add a fresh run with the new text
    p0.add_run(text)
    # Remove any extra paragraphs after the first
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def clone_row_of(table, src_idx=-1):
    """Append a clone of row src_idx to table, returning the new row."""
    src_xml = table.rows[src_idx]._element
    new_xml = deepcopy(src_xml)
    table._element.append(new_xml)
    return table.rows[-1]


def clear_row_text(row):
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in list(p.runs):
                run._element.getparent().remove(run._element)


# ---------- A. Version History ---------------------------------------------

tbl_version = doc.tables[0]
# Row 5 is empty placeholder; fill it
v_row = tbl_version.rows[5]
v_cells = v_row.cells
set_cell_text(v_cells[0], "3")
set_cell_text(v_cells[1], "17/4/2026")
set_cell_text(v_cells[2], "Team Group 1 (final v3)")
set_cell_text(
    v_cells[3],
    "Uplift to BRD Final v3: preserve UC-01..UC-16; add UC-17..UC-26; "
    "align FR/NFR and UC Specifications with locked Business Rules "
    "(Comp, Forfeited, Refund Pending/Refunded, Write-off, shift close, "
    "Reorder drift, 86'd independence); keep Delivery out of scope."
)
print("[A] Version History row 5 populated with v3 Final entry.")


# ---------- B. Glossary additions ------------------------------------------

GLOSSARY_ADDS = [
    ("Reorder",
     "Creating a new order from a previously Completed order; current menu availability and prices apply; removed or changed items surface as inline drift notices before checkout."),
    ("86'd Item",
     "A menu item marked temporarily unavailable by the kitchen; instantly locked on Client WebApp and POS for new orders."),
    ("Refund Pending",
     "Payment status indicating a refund request has been raised and is awaiting processing or Manager confirmation."),
    ("Refunded",
     "Payment status indicating a refund has been completed via gateway callback or Manager manual confirmation."),
    ("Write-off",
     "Manager-approved loss-recording state for a Dine-in order where the customer left without paying; not counted as revenue."),
    ("Comp",
     "Manager-approved voiding of an order (proactive waiver for goodwill or serious error); amount is reclassified as comp expense, not revenue."),
    ("Forfeited",
     "Payment status for a Paid Takeaway/Pickup order where the customer did not arrive within the holding window and the Manager decided the shop retains the payment."),
    ("Shift Close",
     "End-of-shift procedure covering cash drawer reconciliation, payment-method breakdown, and inherited-exception handover to the next shift."),
    ("Audit Log",
     "Immutable record of sensitive financial and operational actions (refund, Comp, Write-off, Forfeited, manager override, shift variance)."),
    ("Table QR Code",
     "Fixed QR code on each dine-in table that encodes the table identifier for QR ordering."),
]

tbl_glossary = doc.tables[4]
for term, definition in GLOSSARY_ADDS:
    new_row = clone_row_of(tbl_glossary)
    clear_row_text(new_row)
    set_cell_text(new_row.cells[0], term)
    set_cell_text(new_row.cells[1], definition)
print(f"[B] Added {len(GLOSSARY_ADDS)} glossary terms.")


# ---------- C. Add A8 scope-lock assumption ---------------------------------

tbl_assumptions = doc.tables[7]
# This table mixes Assumptions and Constraints. Need to insert A8 after A7 (row 7).
# Row indices: 0 header, 1..7 A1..A7, 8 header "Constraints", 9..14 C1..C6.
# Insert new row after row 7 by copying row 7 and inserting.
a7_xml = tbl_assumptions.rows[7]._element
new_a8 = deepcopy(a7_xml)
a7_xml.addnext(new_a8)
# Now rows reload; find new row 8 (previously row 8 Constraint header is now row 9)
# Refresh table reference
tbl_assumptions = doc.tables[7]
a8 = tbl_assumptions.rows[8]
clear_row_text(a8)
set_cell_text(a8.cells[0], "A8")
set_cell_text(
    a8.cells[1],
    "Business scope was locked on 2026-04-17. BRD final v3 covers a curated "
    "26-UC authoring lane (UC-01..UC-16 preserved plus UC-17..UC-26 added); "
    "any new business capability requires a scope-change ADR."
)
print("[C] Added assumption A8 (scope lock) to assumptions table.")


# ---------- D. FR table ----------------------------------------------------

tbl_fr = doc.tables[9]

# Rewrite FR-01 rationale
set_cell_text(
    tbl_fr.rows[1].cells[3],
    "This capability supports the primary customer ordering flow within the BRD final v3 scope."
)
# Rewrite FR-04 rationale
set_cell_text(
    tbl_fr.rows[4].cells[3],
    "This capability supports post-purchase visibility for registered customers within the BRD final v3 scope."
)
# Extend FR-07 rationale
set_cell_text(
    tbl_fr.rows[7].cells[3],
    "This capability provides the reporting data needed for operational and business review, and feeds shift close and end-of-day reconciliation (UC-21)."
)

NEW_FR = [
    ("FR-17", "High",
     "The system shall allow registered customers to reorder items from a Completed past order by applying current menu availability and prices, and by flagging any drift (deleted, archived, repriced, 86'd, or removed-topping items) with inline notices before checkout.",
     "Enables a repeat-purchase workflow consistent with the locked Reorder policy (ADR-025).",
     "UC-17"),
    ("FR-18", "Medium",
     "The system shall allow Managers to create, edit, deactivate, and review percentage and fixed-amount promotion codes, including validity period, usage limits, stacking rules, and a usage audit history.",
     "Supports pricing governance and promotion auditability; free-item and BOGO promotions are out of scope.",
     "UC-18"),
    ("FR-19", "Medium",
     "The system shall provide an operational dashboard for Managers showing revenue snapshot, order counts by status, open exceptions (Refund Pending, overdue Pickup, unresolved Write-off), and shift status.",
     "Enables daily at-a-glance operational oversight separated from detailed reporting.",
     "UC-19"),
    ("FR-20", "High",
     "The system shall allow FOH Staff and Managers to monitor and act on the live active-order queue across Dine-in, Takeaway, and Pickup, including confirmation, cancellation, and exception escalation.",
     "Supports real-time operational control of in-flight orders.",
     "UC-20"),
    ("FR-21", "High",
     "The system shall allow Cashiers and Managers to close a shift with cash-drawer reconciliation, payment-method breakdown (cash, bank QR, gateway, refund, Comp, Forfeited, Write-off), and inherited-exception handover to the next shift.",
     "Supports end-of-day financial accuracy and shift-span attribution per locked shift-close rules.",
     "UC-21"),
    ("FR-22", "High",
     "The system shall allow BOH Staff and Managers to mark a menu item as 86'd and shall instantly propagate availability changes to the Client WebApp and POS, and flag affected not-yet-cooked orders for FOH/Manager action.",
     "Prevents ordering of unavailable items and handles 86'd propagation on in-flight orders.",
     "UC-22"),
    ("FR-23", "High",
     "The system shall allow FOH Service Staff to confirm order handoff and transition Ready orders to Completed, and shall support wrong-handoff correction by reverting Completed to Ready within the same shift with an audit reason.",
     "Closes the dine-in service lifecycle and supports mistaken-completion correction.",
     "UC-23"),
    ("FR-24", "High",
     "The system shall allow FOH Staff and Managers to record customer complaints and operational exceptions (wrong item, missing item, quality issue, wrong handoff) and select a resolution from remake, substitution, goodwill discount, partial refund, full refund, Comp, or Write-off with audit.",
     "Supports service recovery and maintains financial integrity across exception paths.",
     "UC-24"),
    ("FR-25", "Medium",
     "The system shall allow customers to register, log in, log out, update profile, change password, and recover password via OTP.",
     "Supports the authenticated customer lifecycle required for order history, reorder, and rating.",
     "UC-25"),
    ("FR-26", "Critical",
     "The system shall allow customers to complete checkout by confirming a service model, applying an optional promotion code, capturing Guest or Registered identification, reviewing the total, and routing to payment or tracking based on the service model.",
     "Closes the customer-side ordering funnel and enforces service-model-specific prepayment rules.",
     "UC-26"),
]

for req, prio, desc, ration, uc_ref in NEW_FR:
    new_row = clone_row_of(tbl_fr)
    clear_row_text(new_row)
    set_cell_text(new_row.cells[0], req)
    set_cell_text(new_row.cells[1], prio)
    set_cell_text(new_row.cells[2], desc)
    set_cell_text(new_row.cells[3], ration)
    set_cell_text(new_row.cells[4], uc_ref)
print(f"[D] FR table: rationales cleaned and {len(NEW_FR)} new FR rows added (FR-17..FR-26).")


# ---------- E. NFR table ---------------------------------------------------

tbl_nfr = doc.tables[10]

# NFR-02: "Entire 16-UC scope" -> 26-UC BRD final v3 scope
set_cell_text(
    tbl_nfr.rows[2].cells[4],
    "Entire 26-UC BRD final v3 scope (UC-01..UC-26)"
)

# NFR-04: extend UC list to include UC-17..UC-26
set_cell_text(
    tbl_nfr.rows[4].cells[4],
    "UC-01, UC-03, UC-04, UC-05, UC-06, UC-07, UC-08, UC-09, UC-10, UC-11, UC-12, "
    "UC-13, UC-14, UC-15, UC-16, UC-17, UC-18, UC-19, UC-20, UC-21, UC-22, UC-23, "
    "UC-24, UC-25, UC-26"
)

# NFR-05: extend concurrent-user use cases
set_cell_text(
    tbl_nfr.rows[5].cells[4],
    "UC-01, UC-02, UC-09, UC-11, UC-16, UC-20, UC-26"
)

# NFR-07: "UI-facing UCs within the selected 16 UCs" -> 26 UCs of BRD final v3
set_cell_text(
    tbl_nfr.rows[7].cells[4],
    "UI-facing UCs within the 26 UCs of BRD final v3"
)

NEW_NFR = [
    ("NFR-08", "High",
     "The system shall maintain an immutable audit log for sensitive financial and operational actions (refund, Comp, Write-off, Forfeited, manager override, shift-close variance approval) including actor, role, timestamp, related order, affected amount, reason, and approver where applicable.",
     "Supports financial traceability and dispute resolution per the locked Manager Override and Shift Close rules (ADR-022, ADR-023, ADR-024, ADR-026).",
     "UC-10, UC-21, UC-22, UC-24"),
    ("NFR-09", "Medium",
     "The system shall anonymize Guest customer personal data (Name and Phone Number) 12 months after the related order reaches a terminal state (Completed / Cancelled / Refunded / Write-off / Comp / Forfeited), in line with Vietnam's Decree 13/2023/ND-CP on personal data protection.",
     "Supports compliance with personal-data law while preserving order-level statistical data.",
     "UC-01, UC-11, UC-26"),
]

for req, prio, desc, ration, uc_ref in NEW_NFR:
    new_row = clone_row_of(tbl_nfr)
    clear_row_text(new_row)
    set_cell_text(new_row.cells[0], req)
    set_cell_text(new_row.cells[1], prio)
    set_cell_text(new_row.cells[2], desc)
    set_cell_text(new_row.cells[3], ration)
    set_cell_text(new_row.cells[4], uc_ref)
print(f"[E] NFR table: scope references updated; {len(NEW_NFR)} new NFR rows added.")


# ---------- G. Insert UC-17..UC-26 spec tables -----------------------------

UC_SPECS = [
    {
        "id_name": "UC-17: Reorder Past Order",
        "actor": "Registered Customer",
        "description": (
            "The Registered Customer reorders items from a Completed past order. "
            "The system creates a new order using the current menu state: deleted "
            "or archived items are dropped with a notice, repriced items use the "
            "current price, removed toppings are dropped with a notice, and 86'd "
            "items are kept but flagged to block checkout until cleared."
        ),
        "trigger": "The customer selects Reorder on a past order from Order History.",
        "pre": (
            "1. The customer is authenticated (Registered Customer). "
            "2. The source order is in Completed state. "
            "3. The Reorder feature is enabled on the Client WebApp."
        ),
        "post": (
            "Success: A new draft order is created using current menu state, all drift "
            "notices are visible to the customer, and the customer can proceed to "
            "checkout (UC-26). Failure: No new order is committed; the source order "
            "history is unchanged and the system displays an appropriate error."
        ),
        "normal": (
            "1.1. The customer opens Order History. "
            "1.2. The customer selects a past Completed order. "
            "1.3. The customer clicks Reorder. "
            "1.4. The system loads items from the source order and applies current menu state. "
            "1.5. The system displays a consolidated drift banner plus per-line inline notices. "
            "1.6. The customer reviews items and totals. "
            "1.7. The customer accepts and is routed to checkout (UC-26)."
        ),
        "alt": (
            "2a. No drift: The system loads items directly at current prices without notices. "
            "2b. All items unavailable: The system displays \"No items from this order are currently available\" "
            "with a link to the menu. "
            "2c. Partial drop: Only available items are carried over; dropped items are listed in the summary."
        ),
        "exc": (
            "E1. Source order not Completed: Reorder button is hidden. "
            "E2. Item currently 86'd: Kept in cart but blocks checkout until removed or 86'd is cleared. "
            "E3. Promotion code from original order is expired: Removed silently with a notice."
        ),
        "priority": "Medium",
    },
    {
        "id_name": "UC-18: Manage Promotions",
        "actor": "Manager",
        "description": (
            "The Manager creates, edits, deactivates, and reviews promotion codes "
            "(percentage or fixed amount only) including validity period, usage limits, "
            "stacking policy, and usage audit history. Free-item and BOGO promotions "
            "are out of scope."
        ),
        "trigger": "The Manager opens Promotion Management in the Admin WebApp.",
        "pre": (
            "1. The Manager is authenticated with promotion management permission. "
            "2. The Promotion Management module is available."
        ),
        "post": (
            "Success: A promotion code is created, updated, or deactivated with an "
            "audit log entry; usage history remains available for review. "
            "Failure: No promotion change is persisted; the system shows a validation error."
        ),
        "normal": (
            "1.1. The Manager opens Promotion Management. "
            "1.2. The system displays existing promotions with status and current usage. "
            "1.3. The Manager clicks New Promotion. "
            "1.4. The system displays the form (Code, Type, Value, Valid From/To, Conditions, Usage Limit). "
            "1.5. The Manager enters data and saves. "
            "1.6. The system validates and persists the promotion."
        ),
        "alt": (
            "2a. Edit: The Manager selects an existing promotion, edits fields, and saves. "
            "2b. Deactivate: The Manager deactivates a promotion without deleting usage history. "
            "2c. Usage history: The Manager opens the history view for orders that applied a given promotion."
        ),
        "exc": (
            "E1. Duplicate promotion code. "
            "E2. Invalid date range (end before start). "
            "E3. Invalid discount value (negative, zero, or greater than 100% for percentage). "
            "E4. Attempt to configure free-item/BOGO promotion — rejected as out of scope."
        ),
        "priority": "Medium",
    },
    {
        "id_name": "UC-19: View Operational Dashboard",
        "actor": "Manager",
        "description": (
            "The Manager views an operational dashboard summarizing today's revenue, "
            "order counts by status, kitchen load, open exceptions (Refund Pending, "
            "overdue Pickup, unresolved Write-off, Forfeited pending, Dine-in "
            "Completed+Unpaid), and shift status awaiting review."
        ),
        "trigger": "The Manager opens the Admin WebApp home dashboard.",
        "pre": "1. The Manager is authenticated.",
        "post": (
            "Success: The dashboard displays the current operational snapshot with alerts. "
            "Failure: The system displays a loading error and preserves the last valid snapshot."
        ),
        "normal": (
            "1.1. The Manager logs in and lands on the Dashboard. "
            "1.2. The system displays today's revenue, order counts by status, average order value, and alert cards. "
            "1.3. The Manager clicks any alert card to drill down into the related active orders or exception lists."
        ),
        "alt": (
            "2a. Filter by shift or by day. "
            "2b. Drill-down into UC-20 Active Orders from any alert card. "
            "2c. Drill-down into UC-21 Shift Close from shift-status alerts."
        ),
        "exc": (
            "E1. Data loading failure — retry option is offered. "
            "E2. No data yet for the selected day — placeholder message shown."
        ),
        "priority": "Medium",
    },
    {
        "id_name": "UC-20: Manage Active Orders",
        "actor": "Front-of-House Staff / Manager",
        "description": (
            "The FOH Staff or Manager monitors and acts on live orders across Dine-in, "
            "Takeaway, and Pickup — confirming, cancelling (with Manager approval for "
            "paid orders), transferring, and escalating exceptions such as complaints, "
            "wrong handoff, overdue Pickup, and payment recovery."
        ),
        "trigger": "The FOH Staff opens Active Orders in the Admin WebApp.",
        "pre": (
            "1. The FOH Staff or Manager is authenticated. "
            "2. At least one active order exists (not in a terminal state)."
        ),
        "post": (
            "Success: Targeted orders are updated with the correct next state or escalated "
            "to the Manager for exception resolution. Failure: No unintended state change "
            "is applied; the action is rejected with a reason."
        ),
        "normal": (
            "1.1. The Staff opens Active Orders. "
            "1.2. The system lists active orders grouped by service model and status. "
            "1.3. The Staff selects an order and opens details. "
            "1.4. The Staff takes an allowed action (confirm, cancel with reason, transfer table, escalate)."
        ),
        "alt": (
            "2a. Filter by status, service model, or table. "
            "2b. Escalate to Manager for paid cancellations, refund decisions, Comp, or Forfeited. "
            "2c. Review overdue Pickup/Takeaway and route to UC-24 as needed."
        ),
        "exc": (
            "E1. Order already in a terminal state — action disallowed. "
            "E2. Cancellation of a Paid order — must route through refund flow (§3b). "
            "E3. Connectivity loss — local changes are blocked until reconnect."
        ),
        "priority": "High",
    },
    {
        "id_name": "UC-21: Close Shift and Reconcile End-of-day",
        "actor": "Cashier / Manager",
        "description": (
            "The Cashier or Manager closes a shift by verifying the cash drawer count, "
            "reconciling transactions across cash, bank-transfer QR, online gateway, "
            "refund, Write-off, Comp, and Forfeited buckets, and confirming inherited "
            "exceptions are resolved or explicitly handed over to the next shift."
        ),
        "trigger": "The Cashier opens Close Shift at end of shift.",
        "pre": (
            "1. An open shift exists owned by the current Cashier (or Manager acts on their behalf). "
            "2. At least one financial event occurred in the shift, or the shift is to be closed empty."
        ),
        "post": (
            "Success: The shift is closed, variance is explained or accepted, and the shift report "
            "is frozen. Inherited exceptions are listed for the next shift opening. "
            "Failure: The shift remains open; blockers and unresolved items are listed."
        ),
        "normal": (
            "1.1. The Cashier opens Close Shift. "
            "1.2. The system displays expected cash based on paid cash transactions minus cash refunds. "
            "1.3. The Cashier enters actual counted cash. "
            "1.4. The system computes the variance. "
            "1.5. If variance exceeds the configured threshold, Manager confirmation with reason is required. "
            "1.6. The system displays inherited-exception candidates (Refund Pending, Write-off pending, "
            "Forfeited pending, Dine-in Completed+Unpaid). "
            "1.7. The Cashier resolves or marks each as inherited with a note. "
            "1.8. The system closes the shift and locks reports."
        ),
        "alt": (
            "2a. Variance above threshold triggers Manager approval flow. "
            "2b. Manager may reopen a shift with audit reason. "
            "2c. End-of-day reconciliation aggregates all shifts of the day."
        ),
        "exc": (
            "E1. Open Refund Pending blocks close unless inherited with reason. "
            "E2. Hardware or connectivity issue — allow manual cash count with a flag for later review. "
            "E3. Attempt to reopen a reconciled shift without Manager approval — rejected."
        ),
        "priority": "High",
    },
    {
        "id_name": "UC-22: Mark Menu Item as 86'd",
        "actor": "BOH Staff / Manager",
        "description": (
            "The BOH Staff or Manager marks a menu item as 86'd (temporarily unavailable). "
            "The system immediately hides the item from new orders, flags it on existing "
            "not-yet-cooked orders for FOH/Manager action, and supports clearing 86'd when "
            "the item is available again. 86'd is independent from manual stock adjustment."
        ),
        "trigger": "Kitchen detects an ingredient is out of stock or Manager decides to stop selling an item temporarily.",
        "pre": "1. The Staff is authenticated with Kitchen or Manager permission.",
        "post": (
            "Success: The item is flagged 86'd; Client WebApp and POS block new selection; "
            "in-flight orders containing the 86'd item raise an exception flag routed to UC-24. "
            "Failure: No availability state change."
        ),
        "normal": (
            "1.1. The Staff locates the item in the availability panel. "
            "1.2. The Staff clicks Mark as 86'd. "
            "1.3. The system prompts for a short reason. "
            "1.4. The Staff confirms. "
            "1.5. The system propagates 86'd status across Client WebApp and POS and flags affected in-flight orders."
        ),
        "alt": (
            "2a. Clear 86'd when the item is available again. "
            "2b. Apply 86'd to a topping, not the base item. "
            "2c. Bulk 86'd by category (if supported by the config)."
        ),
        "exc": (
            "E1. 86'd does not auto-adjust stock balance (independence rule, §3m). "
            "E2. Paid order containing a now-86'd item — routed to UC-24 complaint/exception handling."
        ),
        "priority": "High",
    },
    {
        "id_name": "UC-23: Serve and Confirm Handoff",
        "actor": "Front-of-House Service Staff / Cashier",
        "description": (
            "The Server confirms actual handoff of a Ready order: for Dine-in, food is "
            "delivered to the correct table; for Takeaway/Pickup, the Cashier hands the "
            "bag to the correct customer. The order transitions from Ready to Completed. "
            "The UC covers wrong-handoff correction by reverting Completed to Ready in "
            "the same shift with audit reason."
        ),
        "trigger": "The kitchen marks the order as Ready.",
        "pre": "1. The order is in Ready state. 2. The handling Staff is authenticated.",
        "post": (
            "Success: The order is set to Completed with the confirming actor logged. "
            "Failure: The order remains Ready and can be flagged for wrong-handoff correction."
        ),
        "normal": (
            "1.1. The Server/Cashier sees the Ready order on their queue. "
            "1.2. The Server delivers food to the correct Dine-in table, or the Cashier "
            "hands the Takeaway/Pickup bag to the matching customer. "
            "1.3. The Server/Cashier taps Delivered or Customer Picked Up. "
            "1.4. The system sets order_status = Completed."
        ),
        "alt": (
            "2a. Wrong handoff correction: Manager or FOH reverts Completed to Ready within the same shift with a required audit reason. "
            "2b. Missing item on arrival: the Server flags the order and routes the issue to UC-24. "
            "2c. Customer arrives earlier than scheduled Pickup: Cashier informs the waiting time; order remains Ready until truly served."
        ),
        "exc": (
            "E1. Customer reports a missing or wrong item on receipt — route to UC-24. "
            "E2. Late Pickup/Takeaway no-show — escalate to Manager; do not auto-Complete. "
            "E3. System connectivity loss — staff uses manual handoff note to reconcile later."
        ),
        "priority": "High",
    },
    {
        "id_name": "UC-24: Handle Complaint and Operational Exception",
        "actor": "Front-of-House Staff / Manager",
        "description": (
            "Staff records a customer complaint or operational exception (wrong item, "
            "missing item, quality issue, wrong handoff) and routes it to a resolution: "
            "remake, substitution, goodwill discount, partial refund, full refund, Comp, "
            "or Write-off. All financial outcomes carry audit logs and flow to shift-close "
            "reconciliation."
        ),
        "trigger": "A customer reports an issue, or Staff notices a service error.",
        "pre": (
            "1. The related order exists and is accessible. "
            "2. The acting Staff has permission; Manager approval is required for refunds, Comp, and Write-off."
        ),
        "post": (
            "Success: The complaint is logged with a resolution and audit trail; financial effects "
            "appear on shift reports. Failure: No resolution is applied; the case remains open for handover."
        ),
        "normal": (
            "1.1. The Staff opens the impacted order. "
            "1.2. The Staff selects Report Issue and picks an issue category. "
            "1.3. The Staff enters a description. "
            "1.4. The system routes the decision path (FOH for minor cases, Manager for financial ones). "
            "1.5. The selected resolution is applied (remake / substitution / refund / Comp / Write-off)."
        ),
        "alt": (
            "2a. Remake does not create new revenue. "
            "2b. Partial refund keeps the order Completed but flags the refund line. "
            "2c. Comp marks the whole order as zero-revenue under Manager authority. "
            "2d. Write-off marks a Dine-in loss where the customer left without paying."
        ),
        "exc": (
            "E1. Refund, Comp, or Write-off without Manager approval — rejected. "
            "E2. Guest PII has been anonymized (12-month retention) — reason is still logged, but PII fields are hidden. "
            "E3. Related order already in a terminal Refunded state — further refund requires a new request and Manager justification."
        ),
        "priority": "High",
    },
    {
        "id_name": "UC-25: Manage Customer Account and Authentication",
        "actor": "Customer (Registered)",
        "description": (
            "The Registered Customer manages account lifecycle: optional registration, "
            "log in, log out, update profile, change password, and recover password via "
            "OTP. Guest Checkout does not require this UC."
        ),
        "trigger": "The customer opens the Account menu on the Client WebApp.",
        "pre": "1. The Client WebApp is available.",
        "post": (
            "Success: The customer's account state reflects the requested change. "
            "Failure: No account change is persisted; the system displays a validation error."
        ),
        "normal": (
            "1.1. The customer opens the Account menu. "
            "1.2. The customer enters phone or email and password. "
            "1.3. The system authenticates. "
            "1.4. The system opens the account home."
        ),
        "alt": (
            "2a. Register: customer enters Name, Phone or Email, and password, and confirms via OTP. "
            "2b. Forgot password: customer requests OTP and sets a new password. "
            "2c. Update profile: customer updates editable fields. "
            "2d. Change password: customer verifies the old password and sets a new one. "
            "2e. Log out."
        ),
        "exc": (
            "E1. Invalid credentials. "
            "E2. OTP expired or wrong. "
            "E3. Password does not meet policy."
        ),
        "priority": "Medium",
    },
    {
        "id_name": "UC-26: Complete Checkout and Apply Promotion",
        "actor": "Customer (Registered / Guest)",
        "description": (
            "At checkout, the customer confirms a service model (Dine-in / Takeaway / "
            "Pickup), applies an optional promotion code, captures identification "
            "(Guest Name + Phone or authenticated profile), reviews the total, and is "
            "routed to payment (Takeaway/Pickup) or to order tracking (Dine-in)."
        ),
        "trigger": "The customer clicks Checkout from the cart.",
        "pre": (
            "1. The cart contains at least one available item. "
            "2. The selected service model is supported."
        ),
        "post": (
            "Success: An order is created with the correct service model, an applied promotion "
            "if any, and identification captured; the customer is routed to payment or tracking. "
            "Failure: No order is created; the cart content is preserved."
        ),
        "normal": (
            "1.1. The customer opens the cart. "
            "1.2. The system displays the cart summary. "
            "1.3. The customer selects a service model. "
            "1.4. The customer enters an optional promotion code. "
            "1.5. The system validates the promotion and updates the total. "
            "1.6. The customer enters Name and Phone (Guest) or confirms an authenticated profile. "
            "1.7. The customer clicks Place Order. "
            "1.8. The system creates the order and routes the customer to UC-02 for prepay or UC-11 for tracking."
        ),
        "alt": (
            "2a. No promotion code. "
            "2b. Invalid promotion — the system removes it and shows a notice. "
            "2c. Guest continues without logging in."
        ),
        "exc": (
            "E1. Empty cart. "
            "E2. Promotion expired or over-used. "
            "E3. Promotion not applicable to the cart or service model."
        ),
        "priority": "Critical",
    },
]

# The last UC spec table is Tables[26] (UC-16). We clone it and insert after it.
uc_template_xml = doc.tables[26]._element
insert_after = uc_template_xml

for spec in UC_SPECS:
    new_tbl_xml = deepcopy(uc_template_xml)
    insert_after.addnext(new_tbl_xml)
    insert_after = new_tbl_xml  # next insert goes right after this one

# Refresh document.tables; new tables are 27..36
doc2 = doc  # same Document instance
new_tables = doc2.tables[27:27 + len(UC_SPECS)]

for tbl, spec in zip(new_tables, UC_SPECS):
    # UC-16 template rows:
    # r0 ID & Name, r1 Primary Actor, r2 Description, r3 Trigger,
    # r4 Pre-conditions, r5 Post-conditions, r6 Normal Flow,
    # r7 Alternative Flows, r8 Exceptions, r9 Priority,
    # r10 Use Case Graphic (if any)
    rows = tbl.rows
    set_cell_text(rows[0].cells[1], spec["id_name"])
    set_cell_text(rows[1].cells[1], spec["actor"])
    set_cell_text(rows[2].cells[1], spec["description"])
    set_cell_text(rows[3].cells[1], spec["trigger"])
    set_cell_text(rows[4].cells[1], spec["pre"])
    set_cell_text(rows[5].cells[1], spec["post"])
    set_cell_text(rows[6].cells[1], spec["normal"])
    set_cell_text(rows[7].cells[1], spec["alt"])
    set_cell_text(rows[8].cells[1], spec["exc"])
    set_cell_text(rows[9].cells[1], spec["priority"])
    # row 10 "Use Case Graphic" left blank
    set_cell_text(rows[10].cells[1], "")

print(f"[G] Inserted {len(UC_SPECS)} UC specification tables (UC-17..UC-26).")


# ---------- Save ----------------------------------------------------------

doc.save(str(SRC))
print()
print(f"Saved {SRC.name}")
