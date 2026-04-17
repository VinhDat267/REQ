from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

root = Path(__file__).resolve().parent.parent
part1_path = root / 'UC_Specifications_Part1_EN.md'
part2_path = root / 'UC_Specifications_Part2_EN.md'
out_path = root / 'UC_Specifications_English_Combined.docx'


def clean_inline(text: str) -> str:
    text = text.replace('\ufeff', '').replace('🔹', '').strip()
    text = text.replace('`', '')
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = text.replace('→', '->')
    text = text.replace('≥', '>=')
    text = text.replace('≤', '<=')
    text = text.replace('–', '-')
    text = text.replace('—', '-')
    return text.strip()


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def set_cell_width(cell, width_twips):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.first_child_found_in('w:tcW')
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_toc(paragraph):
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    text = OxmlElement('w:t')
    text.text = 'Right-click and update field to generate the table of contents.'
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(text)
    run._r.append(fldChar3)


def format_run(run, size=None, bold=None, italic=None, color=None):
    run.font.name = 'Times New Roman'
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def parse_table_rows(rows):
    parsed = []
    for line in rows:
        if line.strip().startswith('|'):
            parsed.append([clean_inline(c.strip()) for c in line.strip().strip('|').split('|')])
    if len(parsed) >= 2 and all(re.fullmatch(r'[-:]+', c.strip()) for c in parsed[1]):
        parsed.pop(1)
    return parsed


def add_md_table(doc, rows):
    parsed = parse_table_rows(rows)
    if not parsed:
        return
    cols = max(len(r) for r in parsed)
    table = doc.add_table(rows=0, cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for r_idx, row in enumerate(parsed):
        row = row + [''] * (cols - len(row))
        cells = table.add_row().cells
        for c_idx, value in enumerate(row):
            cell = cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(value)
            format_run(run, size=10.5)
            if r_idx == 0:
                run.bold = True
                set_cell_shading(cell, 'D9EAF7')
            elif cols == 2 and c_idx == 0:
                run.bold = True
                set_cell_shading(cell, 'F2F2F2')

    if cols == 2:
        for row in table.rows:
            set_cell_width(row.cells[0], 2200)
            set_cell_width(row.cells[1], 7000)

    doc.add_paragraph()


def add_body_paragraph(doc, text, style=None, bold=False, italic=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(clean_inline(text))
    format_run(run, size=11, bold=bold, italic=italic)
    return p


def add_part(doc, part_title, markdown_text):
    doc.add_page_break()
    h = doc.add_paragraph()
    h.style = doc.styles['Heading 1']
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(8)
    run = h.add_run(part_title)
    format_run(run, size=16, bold=True, color='1F4E79')

    lines = markdown_text.splitlines()
    i = 0
    while i < len(lines):
        leading = clean_inline(lines[i].strip())
        if not leading or leading == '---' or leading.startswith('#'):
            i += 1
            continue
        break
    first_member = True
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped or stripped == '---':
            i += 1
            continue
        if stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            add_md_table(doc, table_lines)
            continue
        if stripped.startswith('# '):
            i += 1
            continue
        if stripped.startswith('## '):
            p = doc.add_paragraph()
            p.style = doc.styles['Heading 2']
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(clean_inline(stripped[3:]))
            format_run(run, size=13, bold=True, color='2F75B5')
        elif stripped.startswith('### '):
            if not first_member:
                doc.add_page_break()
            first_member = False
            p = doc.add_paragraph()
            p.style = doc.styles['Heading 2']
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(clean_inline(stripped[4:]))
            format_run(run, size=13, bold=True, color='1F1F1F')
        elif stripped.startswith('#### '):
            p = doc.add_paragraph()
            p.style = doc.styles['Heading 3']
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(clean_inline(stripped[5:]))
            format_run(run, size=12, bold=True, color='C00000')
        elif stripped.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.1
            run = p.add_run(clean_inline(stripped[2:]))
            format_run(run, size=10.5, italic=True, color='666666')
        elif stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(clean_inline(stripped[2:]))
            format_run(run, size=11)
        else:
            add_body_paragraph(doc, stripped)
        i += 1


def build():
    doc = Document()
    props = doc.core_properties
    props.title = 'Wonton POS - Use Case Specifications (English)'
    props.subject = 'Merged midterm use case specifications'
    props.author = 'Claude Code'
    props.company = 'Wonton POS Project Team'
    props.comments = 'Combined and formatted from Part 1 and Part 2 markdown sources.'

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    for style_name in ['Normal', 'Title', 'Subtitle', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4']:
        styles[style_name].font.name = 'Times New Roman'
    styles['Normal'].font.size = Pt(11)
    styles['Heading 1'].font.size = Pt(16)
    styles['Heading 1'].font.bold = True
    styles['Heading 2'].font.size = Pt(13)
    styles['Heading 2'].font.bold = True
    styles['Heading 3'].font.size = Pt(12)
    styles['Heading 3'].font.bold = True

    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_p.add_run('Wonton POS - Use Case Specifications (English)')
    format_run(header_run, size=9, italic=True, color='666666')

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run('Page ')
    format_run(footer_run, size=9, color='666666')
    add_page_number(footer_p)

    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('WONTON POS')
    format_run(r, size=22, bold=True, color='1F4E79')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('USE CASE SPECIFICATIONS')
    format_run(r, size=18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('(English Version)')
    format_run(r, size=14, italic=True, color='666666')

    for _ in range(2):
        doc.add_paragraph()

    for line in [
        'Document Type: Combined midterm deliverable',
        'Source Files: UC_Specifications_Part1_EN.md + UC_Specifications_Part2_EN.md',
        'System Name: Wonton POS',
        'Course Context: Software Requirement Specification (REQ)',
        'Date: 2026-03-11'
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        format_run(run, size=11)

    doc.add_page_break()
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 1']
    run = p.add_run('Table of Contents')
    format_run(run, size=16, bold=True, color='1F4E79')
    add_toc(doc.add_paragraph())

    part1_text = part1_path.read_text(encoding='utf-8')
    part2_text = part2_path.read_text(encoding='utf-8')
    add_part(doc, 'Part 1 - Use Case Assignment & Specifications', part1_text)
    add_part(doc, 'Part 2 - Use Case Specifications', part2_text)

    doc.save(out_path)
    print(f'UPDATED={out_path}')
    print(f'SIZE={out_path.stat().st_size}')


if __name__ == '__main__':
    build()
