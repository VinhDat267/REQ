"""Find references to '16 UC' or scope mentions inside BRD v3 docx."""
import sys, io
from pathlib import Path
from docx import Document

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

doc = Document(str(Path(r"c:/Users/VinhDat/Desktop/REQ/Group1_Tutorial01_BRD_Ver3.docx")))

TARGETS = ["16 UC", "16-UC", "midterm", "Midterm", "16 use case", "sixteen", "Canonical", "Order Status", "Payment Status", "Refund", "Write-off", "Comp", "Forfeited", "scope", "In-Scope", "Out-of-Scope"]

print("=== PARAGRAPH MATCHES ===")
for i, p in enumerate(doc.paragraphs):
    text = p.text or ""
    for t in TARGETS:
        if t.lower() in text.lower():
            short = text.strip().replace("\n", " / ")[:200]
            print(f"[p{i:3d}] {t!r:15s} | {short}")
            break

print()
print("=== TABLE CELL MATCHES ===")
for ti, t in enumerate(doc.tables):
    for ri, row in enumerate(t.rows):
        for ci, cell in enumerate(row.cells):
            text = cell.text or ""
            for tgt in TARGETS:
                if tgt.lower() in text.lower():
                    short = text.strip().replace("\n", " / ")[:200]
                    print(f"[T{ti:2d} r{ri:2d} c{ci}] {tgt!r:15s} | {short}")
                    break
